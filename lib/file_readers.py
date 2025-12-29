"""
生活保護受給者尊厳支援データベース - ファイル読み込みモジュール
Word/Excel/PDF/TXTファイルからのテキスト抽出
"""

import io


def get_supported_extensions() -> dict:
    """サポートするファイル拡張子と説明を返す"""
    return {
        ".docx": "Word文書",
        ".xlsx": "Excelファイル",
        ".pdf": "PDFファイル",
        ".txt": "テキストファイル"
    }


def check_dependencies() -> dict:
    """必要なライブラリの存在チェック"""
    deps = {}
    
    try:
        import python_docx
        deps['python-docx'] = True
    except ImportError:
        deps['python-docx'] = False
    
    try:
        import openpyxl
        deps['openpyxl'] = True
    except ImportError:
        deps['openpyxl'] = False
    
    try:
        import pypdf
        deps['pypdf'] = True
    except ImportError:
        deps['pypdf'] = False
    
    return deps


def read_uploaded_file(uploaded_file) -> str:
    """
    アップロードされたファイルからテキストを抽出
    
    Args:
        uploaded_file: Streamlitのアップロードファイルオブジェクト
        
    Returns:
        抽出されたテキスト
        
    Raises:
        ImportError: 必要なライブラリがない場合
        ValueError: サポートされていないファイル形式の場合
    """
    file_name = uploaded_file.name.lower()
    
    if file_name.endswith('.txt'):
        return _read_txt(uploaded_file)
    elif file_name.endswith('.docx'):
        return _read_docx(uploaded_file)
    elif file_name.endswith('.xlsx'):
        return _read_xlsx(uploaded_file)
    elif file_name.endswith('.pdf'):
        return _read_pdf(uploaded_file)
    else:
        raise ValueError(f"サポートされていないファイル形式です: {file_name}")


def _read_txt(uploaded_file) -> str:
    """テキストファイルを読み込む"""
    content = uploaded_file.read()
    
    # エンコーディングを試行
    for encoding in ['utf-8', 'shift_jis', 'cp932', 'euc-jp']:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    
    # 最終手段
    return content.decode('utf-8', errors='replace')


def _read_docx(uploaded_file) -> str:
    """Word文書を読み込む"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docxがインストールされていません。`uv add python-docx`を実行してください。")
    
    doc = Document(io.BytesIO(uploaded_file.read()))
    
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # テーブルも抽出
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    return "\n".join(paragraphs)


def _read_xlsx(uploaded_file) -> str:
    """Excelファイルを読み込む"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxlがインストールされていません。`uv add openpyxl`を実行してください。")
    
    wb = load_workbook(io.BytesIO(uploaded_file.read()), data_only=True)
    
    all_text = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        all_text.append(f"【シート: {sheet_name}】")
        
        for row in sheet.iter_rows():
            row_values = []
            for cell in row:
                if cell.value is not None:
                    row_values.append(str(cell.value))
            if row_values:
                all_text.append(" | ".join(row_values))
    
    return "\n".join(all_text)


def _read_pdf(uploaded_file) -> str:
    """PDFファイルを読み込む"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdfがインストールされていません。`uv add pypdf`を実行してください。")
    
    reader = PdfReader(io.BytesIO(uploaded_file.read()))
    
    all_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            all_text.append(f"--- ページ {i+1} ---")
            all_text.append(text)
    
    return "\n".join(all_text)
