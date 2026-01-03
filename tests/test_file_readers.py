"""
ファイル読み込みモジュールのテスト
lib/file_readers.py のテスト
"""

import io
import pytest
from unittest.mock import Mock, patch, MagicMock

from lib.file_readers import (
    get_supported_extensions,
    check_dependencies,
    read_uploaded_file,
    _read_txt,
)


class TestGetSupportedExtensions:
    """サポート拡張子取得のテスト"""

    def test_returns_dict(self):
        """辞書を返す"""
        result = get_supported_extensions()
        assert isinstance(result, dict)

    def test_contains_docx(self):
        """docxが含まれる"""
        result = get_supported_extensions()
        assert ".docx" in result
        assert result[".docx"] == "Word文書"

    def test_contains_xlsx(self):
        """xlsxが含まれる"""
        result = get_supported_extensions()
        assert ".xlsx" in result
        assert result[".xlsx"] == "Excelファイル"

    def test_contains_pdf(self):
        """pdfが含まれる"""
        result = get_supported_extensions()
        assert ".pdf" in result
        assert result[".pdf"] == "PDFファイル"

    def test_contains_txt(self):
        """txtが含まれる"""
        result = get_supported_extensions()
        assert ".txt" in result
        assert result[".txt"] == "テキストファイル"

    def test_four_extensions(self):
        """4種類の拡張子"""
        result = get_supported_extensions()
        assert len(result) == 4


class TestCheckDependencies:
    """依存関係チェックのテスト"""

    def test_returns_dict(self):
        """辞書を返す"""
        result = check_dependencies()
        assert isinstance(result, dict)

    def test_contains_python_docx(self):
        """python-docxキーが含まれる"""
        result = check_dependencies()
        assert 'python-docx' in result
        assert isinstance(result['python-docx'], bool)

    def test_contains_openpyxl(self):
        """openpyxlキーが含まれる"""
        result = check_dependencies()
        assert 'openpyxl' in result
        assert isinstance(result['openpyxl'], bool)

    def test_contains_pypdf(self):
        """pypdfキーが含まれる"""
        result = check_dependencies()
        assert 'pypdf' in result
        assert isinstance(result['pypdf'], bool)


class TestReadTxt:
    """テキストファイル読み込みのテスト"""

    def test_read_utf8(self):
        """UTF-8テキストの読み込み"""
        content = "こんにちは世界"
        mock_file = Mock()
        mock_file.read.return_value = content.encode('utf-8')

        result = _read_txt(mock_file)
        assert result == content

    def test_read_shift_jis(self):
        """Shift-JISテキストの読み込み"""
        content = "こんにちは世界"
        mock_file = Mock()
        mock_file.read.return_value = content.encode('shift_jis')

        result = _read_txt(mock_file)
        assert result == content

    def test_read_cp932(self):
        """CP932テキストの読み込み"""
        content = "テスト文字列"
        mock_file = Mock()
        mock_file.read.return_value = content.encode('cp932')

        result = _read_txt(mock_file)
        assert result == content

    def test_read_ascii(self):
        """ASCIIテキストの読み込み"""
        content = "Hello World"
        mock_file = Mock()
        mock_file.read.return_value = content.encode('ascii')

        result = _read_txt(mock_file)
        assert result == content

    def test_read_empty_file(self):
        """空ファイルの読み込み"""
        mock_file = Mock()
        mock_file.read.return_value = b""

        result = _read_txt(mock_file)
        assert result == ""


class TestReadUploadedFile:
    """アップロードファイル読み込みのテスト"""

    def test_read_txt_file(self):
        """txtファイルの読み込み"""
        content = "テストテキスト"
        mock_file = Mock()
        mock_file.name = "test.txt"
        mock_file.read.return_value = content.encode('utf-8')

        result = read_uploaded_file(mock_file)
        assert result == content

    def test_read_txt_uppercase(self):
        """大文字拡張子も処理"""
        content = "テストテキスト"
        mock_file = Mock()
        mock_file.name = "TEST.TXT"
        mock_file.read.return_value = content.encode('utf-8')

        result = read_uploaded_file(mock_file)
        assert result == content

    def test_unsupported_extension(self):
        """サポートされていない拡張子"""
        mock_file = Mock()
        mock_file.name = "test.xyz"

        with pytest.raises(ValueError) as exc_info:
            read_uploaded_file(mock_file)

        assert "サポートされていない" in str(exc_info.value)

    @patch('lib.file_readers._read_docx')
    def test_read_docx_file(self, mock_read_docx):
        """docxファイルの読み込み"""
        mock_read_docx.return_value = "Word文書の内容"
        mock_file = Mock()
        mock_file.name = "test.docx"

        result = read_uploaded_file(mock_file)

        assert result == "Word文書の内容"
        mock_read_docx.assert_called_once_with(mock_file)

    @patch('lib.file_readers._read_xlsx')
    def test_read_xlsx_file(self, mock_read_xlsx):
        """xlsxファイルの読み込み"""
        mock_read_xlsx.return_value = "Excelの内容"
        mock_file = Mock()
        mock_file.name = "data.xlsx"

        result = read_uploaded_file(mock_file)

        assert result == "Excelの内容"
        mock_read_xlsx.assert_called_once_with(mock_file)

    @patch('lib.file_readers._read_pdf')
    def test_read_pdf_file(self, mock_read_pdf):
        """pdfファイルの読み込み"""
        mock_read_pdf.return_value = "PDFの内容"
        mock_file = Mock()
        mock_file.name = "document.pdf"

        result = read_uploaded_file(mock_file)

        assert result == "PDFの内容"
        mock_read_pdf.assert_called_once_with(mock_file)


class TestReadTxtFallback:
    """テキストファイル読み込みのフォールバックテスト"""

    def test_fallback_to_replace_errors(self):
        """不正なバイト列の場合はerrors='replace'で処理"""
        # すべてのエンコーディングで失敗するバイト列
        invalid_bytes = bytes([0x80, 0x81, 0x82, 0xFF, 0xFE])
        mock_file = Mock()
        mock_file.read.return_value = invalid_bytes

        result = _read_txt(mock_file)

        # 結果が文字列であることを確認（エラーなし）
        assert isinstance(result, str)
        # 置換文字が含まれる
        assert len(result) > 0

    def test_euc_jp_encoding(self):
        """EUC-JPテキストの読み込み"""
        content = "日本語テスト"
        mock_file = Mock()
        mock_file.read.return_value = content.encode('euc-jp')

        result = _read_txt(mock_file)
        assert result == content


class TestReadDocx:
    """Word文書読み込みのテスト"""

    def test_docx_basic_structure(self):
        """Word文書読み込みの基本構造確認"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docxがインストールされていません")

        from docx import Document
        import io

        # 実際のWord文書を作成
        doc = Document()
        doc.add_paragraph("テスト段落1")
        doc.add_paragraph("テスト段落2")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_docx
        result = _read_docx(mock_file)

        assert "テスト段落1" in result
        assert "テスト段落2" in result

    def test_docx_with_table(self):
        """テーブルを含むWord文書の読み込み"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docxがインストールされていません")

        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("ヘッダー段落")

        # テーブルを追加
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "セル1"
        table.cell(0, 1).text = "セル2"
        table.cell(1, 0).text = "セル3"
        table.cell(1, 1).text = "セル4"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_docx
        result = _read_docx(mock_file)

        assert "ヘッダー段落" in result
        assert "セル1" in result
        assert "|" in result  # テーブルセル区切り

    def test_docx_empty_paragraphs(self):
        """空の段落は除外される"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docxがインストールされていません")

        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("有効な段落")
        doc.add_paragraph("")  # 空の段落
        doc.add_paragraph("   ")  # 空白のみの段落
        doc.add_paragraph("別の有効な段落")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_docx
        result = _read_docx(mock_file)

        assert "有効な段落" in result
        assert "別の有効な段落" in result


class TestReadXlsx:
    """Excel読み込みのテスト"""

    def test_xlsx_basic_structure(self):
        """Excel読み込みの基本構造確認"""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxlがインストールされていません")

        from openpyxl import Workbook
        import io

        wb = Workbook()
        ws = wb.active
        ws.title = "TestSheet"
        ws['A1'] = "テストデータ"
        ws['B1'] = "2番目のセル"

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_xlsx
        result = _read_xlsx(mock_file)

        assert "TestSheet" in result
        assert "テストデータ" in result

    def test_xlsx_multiple_sheets(self):
        """複数シートのExcel読み込み"""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxlがインストールされていません")

        from openpyxl import Workbook
        import io

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1['A1'] = "シート1のデータ"

        ws2 = wb.create_sheet("Sheet2")
        ws2['A1'] = "シート2のデータ"

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_xlsx
        result = _read_xlsx(mock_file)

        assert "Sheet1" in result
        assert "シート1のデータ" in result
        assert "Sheet2" in result
        assert "シート2のデータ" in result

    def test_xlsx_with_empty_cells(self):
        """空のセルがあるExcelの読み込み"""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxlがインストールされていません")

        from openpyxl import Workbook
        import io

        wb = Workbook()
        ws = wb.active
        ws['A1'] = "データ1"
        ws['C1'] = "データ2"  # B1は空

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_xlsx
        result = _read_xlsx(mock_file)

        assert "データ1" in result
        assert "データ2" in result


class TestReadPdf:
    """PDF読み込みのテスト"""

    def test_pdf_basic_structure(self):
        """PDF読み込みの基本構造確認"""
        try:
            from pypdf import PdfReader
        except ImportError:
            pytest.skip("pypdfがインストールされていません")

        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
        except ImportError:
            pytest.skip("reportlabがインストールされていません")

        import io

        # ReportLabでPDFを作成
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        c.drawString(100, 750, "Test PDF Content")
        c.showPage()
        c.save()
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_pdf
        result = _read_pdf(mock_file)

        # PDFが読み込まれることを確認
        assert isinstance(result, str)
        # ページ番号が含まれる
        assert "ページ 1" in result

    def test_pdf_multiple_pages(self):
        """複数ページPDFの読み込み"""
        try:
            from pypdf import PdfReader
        except ImportError:
            pytest.skip("pypdfがインストールされていません")

        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
        except ImportError:
            pytest.skip("reportlabがインストールされていません")

        import io

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        c.drawString(100, 750, "Page 1 Content")
        c.showPage()
        c.drawString(100, 750, "Page 2 Content")
        c.showPage()
        c.save()
        buffer.seek(0)

        mock_file = Mock()
        mock_file.read.return_value = buffer.getvalue()

        from lib.file_readers import _read_pdf
        result = _read_pdf(mock_file)

        assert "ページ 1" in result
        assert "ページ 2" in result
